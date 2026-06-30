"""
Loaders for different file types.
Each loader returns a Document object.
"""
import os
from pathlib import Path
from models import Document


def load_text(path: str) -> Document:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()
    return Document(content=content, source=path, doc_type="text")


def load_pdf(path: str) -> Document:
    from pdfminer.high_level import extract_text
    content = extract_text(path)
    return Document(content=content, source=path, doc_type="pdf")


def load_docx(path: str) -> Document:
    from docx import Document as DocxDocument
    doc = DocxDocument(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    content = "\n".join(paragraphs)
    return Document(content=content, source=path, doc_type="text")


def load_code(path: str) -> Document:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()
    ext = Path(path).suffix
    return Document(
        content=f"```{ext.lstrip('.')}\n{content}\n```",
        source=path,
        doc_type="code",
        metadata={"language": ext.lstrip(".")},
    )


EXTENSION_MAP = {
    ".txt": load_text,
    ".md": load_text,
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".py": load_code,
    ".js": load_code,
    ".ts": load_code,
    ".java": load_code,
    ".go": load_code,
}


def load_file(path: str) -> Document:
    ext = Path(path).suffix.lower()
    loader = EXTENSION_MAP.get(ext, load_text)
    return loader(path)


def load_directory(directory: str) -> list[Document]:
    docs = []
    for root, _, files in os.walk(directory):
        for f in files:
            full_path = os.path.join(root, f)
            try:
                docs.append(load_file(full_path))
            except Exception as e:
                print(f"[Loader] Skipped {full_path}: {e}")
    return docs
