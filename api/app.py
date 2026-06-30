"""
FastAPI + Gradio application — full port from mini_rag.py.

Endpoints (matching mini_rag.py):
  POST /ingest              — ingest text or files
  POST /query               — standard JSON response with timings
  POST /query/stream        — Server-Sent Events token streaming
  POST /evaluate            — batch embedding-based evaluation
  DELETE /source/{name}     — remove a source and its vectors
  GET  /health              — detailed health check

Gradio UI (matching mini_rag.py):
  Tab 1: Chat — streaming, per-session state, timings footer, source citations
  Tab 2: Add Knowledge — multi-file upload, PDF/DOCX/TXT/MD/JSON/CSV support

RAGAS langchain compat patch applied at top (matching mini_rag.py).
"""

# ─── RAGAS PATCH FOR LANGCHAIN 0.4+ COMPATIBILITY ────────────────────────────
# Matches mini_rag.py's monkey-patch exactly
import sys
import types
try:
    import langchain_community.chat_models
    if not hasattr(langchain_community.chat_models, "vertexai"):
        dummy = types.ModuleType("langchain_community.chat_models.vertexai")
        dummy.ChatVertexAI = type("ChatVertexAI", (object,), {})
        sys.modules["langchain_community.chat_models.vertexai"] = dummy
        langchain_community.chat_models.vertexai = dummy
except Exception:
    pass
# ─────────────────────────────────────────────────────────────────────────────

import asyncio
import concurrent.futures as _cf
import json
import os
import time
import uuid

from fastapi import FastAPI, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from dotenv import load_dotenv

load_dotenv()

from config import get_settings
from models import QueryRequest, QueryResponse
from ingestion.loaders import load_directory, load_file
from ingestion.chunker import StructureAwareChunker
from ingestion.metadata import MetadataEnricher
from retrieval.embedder import Embedder
from retrieval.vector_store import VectorStore
from retrieval.bm25_store import BM25Store
from retrieval.hybrid_retriever import HybridRetriever
from retrieval.query_rewriter import QueryRewriter
from reranking.reranker import Reranker, ContextCompressor
from memory.conversation import ConversationMemory
from evaluation.evaluator import EmbeddingEvaluator
from agent.graph import build_simple_graph

settings = get_settings()

# ─── Singletons (created once, shared across requests) ───────────────────────
embedder       = Embedder()
vector_store   = VectorStore(embedder)
bm25_store     = BM25Store()
query_rewriter = QueryRewriter()
retriever      = HybridRetriever(vector_store, bm25_store, embedder, query_rewriter)
reranker       = Reranker()
compressor     = ContextCompressor()
memory         = ConversationMemory()
evaluator      = EmbeddingEvaluator(embedder=embedder)

# Rebuild BM25 from chunks restored by VectorStore on startup
bm25_store.build(vector_store.all_chunks)

# Build LangGraph + helpers
rag_graph, run, generate, generate_stream, cached_retrieve_and_rerank = (
    build_simple_graph(retriever, reranker, memory)
)

# ─── Request models ───────────────────────────────────────────────────────────

class IngestReq(BaseModel):
    text: str
    source: str = "manual"

class QueryReq(BaseModel):
    query: str
    session_id: str = "default"

class EvalReq(BaseModel):
    questions: list[str]
    ground_truths: list[str] = []

# ─── FastAPI app ──────────────────────────────────────────────────────────────
app = FastAPI(title="Mini RAG")

# ─── /ingest ─────────────────────────────────────────────────────────────────
@app.post("/ingest")
def ingest(req: IngestReq):
    """
    Ingest raw text. Matches mini_rag.py's /ingest endpoint exactly:
      - 409 on duplicate source
      - Chunks + metadata enrichment + embedding + BM25 rebuild
    """
    if req.source in vector_store.ingested_sources:
        raise HTTPException(
            status_code=409,
            detail=f"Source '{req.source}' is already indexed. Delete it first or use a unique source name.",
        )
    try:
        chunker  = StructureAwareChunker()
        enricher = MetadataEnricher()

        # Build a Document-like object for the chunker
        from models import Document
        doc    = Document(content=req.text, source=req.source)
        chunks = chunker.chunk(doc)
        chunks = enricher.enrich_batch(chunks)

        vector_store.add(chunks)
        bm25_store.build(vector_store.all_chunks)
        cached_retrieve_and_rerank.cache_clear()

        return {
            "status":         "ok",
            "source":         req.source,
            "chunks_created": len(chunks),
            "total_chunks":   vector_store.total_chunks,
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ─── /query ──────────────────────────────────────────────────────────────────
@app.post("/query")
def query_endpoint(req: QueryReq):
    """Standard blocking query — returns full JSON response with timings."""
    try:
        return run(req.query, req.session_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ─── /query/stream ───────────────────────────────────────────────────────────
@app.post("/query/stream")
async def query_stream(req: QueryReq):
    """
    True token streaming over SSE.
    Matches mini_rag.py's /query/stream exactly:
      - Retrieval in thread pool (non-blocking for the event loop)
      - LLM streaming bridged via asyncio.Queue so sync generator
        never blocks the uvicorn event loop
    """
    loop = asyncio.get_running_loop()

    def _retrieve_only(query: str, session_id: str):
        history   = memory.get_history(session_id)
        rewritten = query if not history.strip() else query_rewriter.rewrite(query, history)
        hits      = list(cached_retrieve_and_rerank(rewritten))
        return rewritten, hits

    rewritten_q, hits = await loop.run_in_executor(
        None, _retrieve_only, req.query, req.session_id
    )

    sources = [
        {
            "id":      i + 1,
            "source":  h.chunk.source,
            "heading": h.chunk.heading,
            "score":   round(h.rerank_score if settings.use_reranker else h.rrf_score, 3),
            "content": h.chunk.content,
        }
        for i, h in enumerate(hits)
    ]

    q: asyncio.Queue = asyncio.Queue(maxsize=32)
    _DONE = object()

    def _stream_to_queue():
        try:
            for token in generate_stream(req.query, rewritten_q, hits):
                loop.call_soon_threadsafe(q.put_nowait, token)
        finally:
            loop.call_soon_threadsafe(q.put_nowait, _DONE)

    full_answer_parts: list[str] = []

    async def gen():
        loop.run_in_executor(None, _stream_to_queue)
        while True:
            item = await q.get()
            if item is _DONE:
                break
            full_answer_parts.append(item)
            yield f"data: {json.dumps({'token': item})}\n\n"
        full_answer = "".join(full_answer_parts).strip()
        memory.add(req.session_id, "user",      req.query)
        memory.add(req.session_id, "assistant", full_answer)
        yield f"data: {json.dumps({'sources': sources, 'done': True})}\n\n"

    return StreamingResponse(gen(), media_type="text/event-stream")


# ─── /evaluate ───────────────────────────────────────────────────────────────
@app.post("/evaluate")
def evaluate_endpoint(req: EvalReq):
    """Batch embedding-based evaluation. Matches mini_rag.py's /evaluate."""
    try:
        return evaluator.evaluate_rag(
            run_fn=lambda q: run(q, session_id=f"eval_{uuid.uuid4()}"),
            questions=req.questions,
            ground_truths=req.ground_truths or None,
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ─── DELETE /source/{name} ───────────────────────────────────────────────────
@app.delete("/source/{source_name}")
def delete_source(source_name: str):
    """Remove all vectors + in-memory chunks for a source. Matches mini_rag.py."""
    removed = vector_store.delete_source(source_name)
    bm25_store.build(vector_store.all_chunks)
    cached_retrieve_and_rerank.cache_clear()
    return {"status": "ok", "source": source_name, "chunks_removed": removed}


# ─── /ingest/file ─────────────────────────────────────────────────────────────
@app.post("/ingest/path")
def ingest_path(path: str, save_index: bool = True):
    """Ingest from a filesystem path (file or directory)."""
    if not os.path.exists(path):
        raise HTTPException(status_code=400, detail=f"Path not found: {path}")
    if os.path.isdir(path):
        docs = load_directory(path)
    else:
        docs = [load_file(path)]

    chunker  = StructureAwareChunker()
    enricher = MetadataEnricher()
    all_chunks = []
    for doc in docs:
        if doc.source in vector_store.ingested_sources:
            continue
        chunks = chunker.chunk(doc)
        chunks = enricher.enrich_batch(chunks)
        all_chunks.extend(chunks)

    if all_chunks:
        vector_store.add(all_chunks)
        bm25_store.build(vector_store.all_chunks)
        cached_retrieve_and_rerank.cache_clear()

    return {"status": "ok", "chunks_created": len(all_chunks), "total": vector_store.total_chunks}


# ─── /health ─────────────────────────────────────────────────────────────────
@app.get("/health")
def health():
    """Detailed health check — matches mini_rag.py's /health."""
    info = vector_store.get_collection_info()
    return {
        "status":           "ok",
        "chunks_in_memory": vector_store.total_chunks,
        "vectors_in_qdrant": info.points_count,
        "active_sessions":  memory.active_sessions,
        "indexed_sources":  list(vector_store.ingested_sources),
    }


# ─── GRADIO UI ───────────────────────────────────────────────────────────────
import gradio as gr

# ─── File text extractor — matches mini_rag.py's extract_text_from_file() ────
def extract_text_from_file(file_path: str) -> tuple[str, str]:
    """Extract text content and derive a source name from an uploaded file."""
    filename    = os.path.basename(file_path)
    source_name = filename
    ext         = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    try:
        if ext in ("txt", "md", "rst", "log", "csv"):
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read(), source_name

        elif ext == "pdf":
            try:
                import pypdf
                reader = pypdf.PdfReader(file_path)
                text   = "\n".join(page.extract_text() or "" for page in reader.pages)
                return text, source_name
            except ImportError:
                return "", "__error__:Install pypdf: pip install pypdf"

        elif ext == "docx":
            try:
                import docx
                doc  = docx.Document(file_path)
                text = "\n".join(p.text for p in doc.paragraphs)
                return text, source_name
            except ImportError:
                return "", "__error__:Install python-docx: pip install python-docx"

        elif ext == "json":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                import json as _json
                data = _json.load(f)
                return _json.dumps(data, indent=2), source_name

        else:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read(), source_name

    except Exception as e:
        return "", f"__error__:{e}"


def ingest_file(files) -> str:
    """Handle one or more uploaded files. Matches mini_rag.py's ingest_file()."""
    if not files:
        return "⚠️ Please upload at least one file."

    results = []
    for file in (files if isinstance(files, list) else [files]):
        file_path   = file.name if hasattr(file, "name") else str(file)
        text, source_name = extract_text_from_file(file_path)

        if source_name.startswith("__error__:"):
            results.append(f"Error: {os.path.basename(file_path)}: {source_name[10:]}")
            continue

        if not text.strip():
            results.append(f"⚠️ {source_name}: No text could be extracted.")
            continue

        if source_name in vector_store.ingested_sources:
            results.append(f"⚠️ '{source_name}' is already indexed. Delete it first or rename the file.")
            continue

        try:
            from models import Document
            doc    = Document(content=text, source=source_name)
            chunker  = StructureAwareChunker()
            enricher = MetadataEnricher()
            chunks = chunker.chunk(doc)
            chunks = enricher.enrich_batch(chunks)
            vector_store.add(chunks)
            results.append(f"✅ Indexed '{source_name}' → {len(chunks)} chunks (total: {vector_store.total_chunks})")
        except Exception as e:
            results.append(f"❌ Error '{source_name}': {e}")

    # Rebuild BM25 once after all files are processed (PERF 3 fix from mini_rag.py)
    bm25_store.build(vector_store.all_chunks)
    cached_retrieve_and_rerank.cache_clear()
    return "\n".join(results)


# ─── Background RAGAS evaluator pool ─────────────────────────────────────────
_ragas_pool = _cf.ThreadPoolExecutor(max_workers=1, thread_name_prefix="ragas")

def _run_ragas_background(question: str, answer: str, contexts: list[str]):
    """Fire-and-forget RAGAS. Matches mini_rag.py's _run_ragas_background()."""
    try:
        res = evaluator.evaluate_single(question=question, answer=answer, contexts=contexts)
        if res.faithfulness is not None:
            parts = []
            if res.faithfulness     is not None: parts.append(f"Faithfulness: {res.faithfulness:.2f}")
            if res.answer_relevancy is not None: parts.append(f"Answer Relevancy: {res.answer_relevancy:.2f}")
            print(f"[RAGAS] {' | '.join(parts)}")
    except Exception as e:
        print(f"[RAGAS] Error: {e}")


def chat_fn(message: str, history: list, session_id: str):
    """
    Streaming chat with background RAGAS evaluation.
    Matches mini_rag.py's chat_fn() exactly:
      - Per-session session_id from gr.State
      - Parallel rewrite + prefetch when history exists
      - Streaming tokens live
      - Timings footer + rewritten query display + source citations
      - Background RAGAS evaluation
    """
    try:
        t0          = time.time()
        history_str = memory.get_history(session_id)

        if not history_str.strip():
            rewritten_q = message
            rewrite_ms  = 0.0
        else:
            t_rw        = time.time()
            rewritten_q = query_rewriter.rewrite(message, history_str)
            rewrite_ms  = round(time.time() - t_rw, 2)

        t_ret  = time.time()
        hits   = list(cached_retrieve_and_rerank(rewritten_q))
        ret_ms = round(time.time() - t_ret, 3)

        sources = [
            {
                "id":      i + 1,
                "source":  h.chunk.source,
                "heading": h.chunk.heading,
                "score":   round(h.rerank_score if settings.use_reranker else h.rrf_score, 3),
                "content": h.chunk.content,
            }
            for i, h in enumerate(hits)
        ]

        answer_parts: list[str] = []
        partial   = ""
        ttft_ms   = None
        t_stream  = time.time()

        for token in generate_stream(message, rewritten_q, hits):
            if ttft_ms is None:
                ttft_ms = round(time.time() - t_stream, 3)
            answer_parts.append(token)
            partial = "".join(answer_parts)
            yield partial

        stream_total_ms = round(time.time() - t_stream, 3)
        if ttft_ms is None:
            ttft_ms = 0.0

        answer = partial.strip()
        gen_ms = round(time.time() - t0, 2)

        if not answer:
            yield "The model returned an empty response. Please try again."
            return

        memory.add(session_id, "user",      message)
        memory.add(session_id, "assistant", answer)

        # ── Footer: timings + rewritten query + sources ──────────────────────
        footer = ""
        timing_parts = []
        if rewrite_ms:  timing_parts.append(f"Rewrite: {rewrite_ms}s")
        timing_parts.append(f"Retrieve: {ret_ms}s")
        timing_parts.append(f"TTFT: {ttft_ms}s")
        timing_parts.append(f"Generation: {round(stream_total_ms - ttft_ms, 3)}s")
        timing_parts.append(f"Total: {gen_ms}s")
        footer += f"\n\n⏱️ **Timings:** `{'  |  '.join(timing_parts)}`"

        if rewritten_q and rewritten_q != message:
            footer += f"\n\n🔍 **Rewritten Query:** `{rewritten_q}`"

        if sources:
            footer += "\n\n**Sources Used:**"
            for s in sources:
                footer += f"\n- [{s['id']}] `{s['source']}` (Confidence: {s['score']})"

        yield answer + footer

        # ── Background RAGAS ─────────────────────────────────────────────────
        contexts = [s["content"] for s in sources]
        if contexts and "blocked by" not in answer.lower():
            _ragas_pool.submit(_run_ragas_background, message, answer, contexts)

    except Exception as e:
        yield f"❌ Error: {str(e)}"


# ─── Gradio UI blocks — matches mini_rag.py's gr.Blocks() exactly ────────────
with gr.Blocks() as demo:
    gr.Markdown("# Mini RAG System")
    gr.Markdown(
        "Production RAG: LangGraph + BM25 + Qdrant Hybrid Search + "
        "Cross-Encoder Reranking (TinyBERT) + Groq Inference."
    )
    # Each browser tab gets its own session ID (BUG 4 fix from mini_rag.py)
    session_state = gr.State(lambda: f"gradio_{uuid.uuid4().hex[:8]}")

    with gr.Tabs():
        with gr.TabItem("Chat"):
            gr.ChatInterface(
                fn=chat_fn,
                additional_inputs=[session_state],
                chatbot=gr.Chatbot(height=500),
                textbox=gr.Textbox(
                    placeholder="Ask a question about the indexed documents...",
                    container=False,
                    scale=7,
                ),
            )
        with gr.TabItem("Add Knowledge"):
            gr.Markdown(
                "### Drop your files below to index them\n"
                "Supported formats: **PDF, DOCX, TXT, MD, CSV, JSON** "
                "(and most plain-text formats).  \n"
                "The filename is automatically used as the source name."
            )
            file_drop = gr.File(
                label="Drop files here or click to upload",
                file_count="multiple",
                type="filepath",
            )
            ingest_btn    = gr.Button("Ingest Files", variant="primary")
            ingest_status = gr.Textbox(label="Ingestion Status", interactive=False, lines=6)
            ingest_btn.click(fn=ingest_file, inputs=[file_drop], outputs=ingest_status)

# Mount Gradio onto the FastAPI app (matches mini_rag.py)
app = gr.mount_gradio_app(app, demo, path="/")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
